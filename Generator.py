# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""

import lxml.etree as ET
import os 

import docx
from lxml import etree

NAME_SPACE = '{http://xmlns.tibco.com/bw/process/2003}'
#Documentation File
MY_DOC = docx.Document()

def listDirectoryContent(dirPathP):
    fileList = []
    for dirpath, dirnames, files in os.walk(dirPathP):
        # print('Directory: ', dirpath)
        for file_name in files:
            if file_name.__contains__('.process'): 
                fileList.append(dirpath + '/'+ file_name)
   
    return fileList




def createFile(content):
    #Create file documentation
    # f = open("documentation.doc", "w")
    # f.write(content)
    
    # f.close()
    
    mydoc = docx.Document()
    mydoc.add_paragraph(content)
    mydoc.save("documentation.docx")
    
def loadFiles():
    dirPathP ='C:/Users/alma_/Desktop/repos/tibco-agents/ApiHebOrderService/'
    listFiles = listDirectoryContent(dirPathP)
    
    content = ''
    
    for x in listFiles:
        print(x)
        getContentFile(x)

    #createFile(content)
    MY_DOC.save("documentation.docx")


def getContentFile(file_name):
    
    recovering_parser = etree.XMLParser(recover=True)
    
    dom = ET.parse("C:\\Users\\alma_\\Desktop\\repos\\tibco-agents\\ApiHebOrderService\\ApiHebOrder Processes\\Main Process\\Send_Order_Main.process", parser=recovering_parser)
    xslt = ET.parse("C:\\alma\\HEB\\Documentator\\Template2.xsl")
    transform = ET.XSLT(xslt)
    newdom = transform(dom)
     
    print(newdom)
 
    data = ET.tostring(newdom)
    outfile = open("filename.xml", 'w')
    outfile.write(data.decode('utf-8'))
    outfile.close()
                    

#unit tests
dirPathP ='C:/Users/alma_/Desktop/repos/tibco-agents/ApiHebOrderService/ApiHebOrder Processes/Sub Process/OMSToWMS.process'
# listT = listDirectoryContent(dirPathP)
 
getContentFile(dirPathP)
# print(listT)
# loadFiles()